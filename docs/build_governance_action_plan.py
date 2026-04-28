from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


DOC_PATH = Path(r"C:\Users\Guilh\AppData\Local\Temp\INOVE_QUATAI\docs\Plano_de_Acao_Governanca_INOVE_QUATAI.docx")

NAVY = RGBColor(18, 59, 101)
SLATE = RGBColor(71, 85, 105)
TEAL = RGBColor(13, 148, 136)
DANGER = RGBColor(185, 28, 28)
AMBER = RGBColor(180, 83, 9)
SOFT_BLUE = "EAF2FB"
SOFT_GRAY = "F8FAFC"
SOFT_RED = "FDECEC"
SOFT_AMBER = "FEF3C7"
SOFT_GREEN = "EAF7EF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=90, start=100, bottom=90, end=100):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="D7E0EA", size=6):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(size))
        element.set(qn("w:color"), color)


def style_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.82)
    section.right_margin = Inches(0.82)

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.2)
    normal.font.color.rgb = RGBColor(31, 41, 55)

    title = doc.styles["Title"]
    title.font.name = "Aptos Display"
    title.font.size = Pt(24)
    title.font.bold = True
    title.font.color.rgb = NAVY

    subtitle = doc.styles["Subtitle"]
    subtitle.font.name = "Aptos"
    subtitle.font.size = Pt(11)
    subtitle.font.italic = False
    subtitle.font.color.rgb = SLATE

    for name, size, color in (
        ("Heading 1", 15, NAVY),
        ("Heading 2", 11.8, TEAL),
        ("Heading 3", 10.5, SLATE),
    ):
        style = doc.styles[name]
        style.font.name = "Aptos"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color

    if "Body Small" not in doc.styles:
        body_small = doc.styles.add_style("Body Small", WD_STYLE_TYPE.PARAGRAPH)
        body_small.base_style = normal
        body_small.font.name = "Aptos"
        body_small.font.size = Pt(9.2)

    if "Callout" not in doc.styles:
        callout = doc.styles.add_style("Callout", WD_STYLE_TYPE.PARAGRAPH)
        callout.base_style = normal
        callout.font.name = "Aptos"
        callout.font.size = Pt(9.8)
        callout.font.color.rgb = RGBColor(30, 41, 59)


def add_para(doc, text, style=None, align=None, before=0, after=6, bold=False, color=None, size=None):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    run = p.add_run(text)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    if size is not None:
        run.font.size = Pt(size)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_cover(doc):
    add_para(doc, "Plano de Acao", style="Title", align=WD_ALIGN_PARAGRAPH.CENTER, before=68, after=2)
    add_para(
        doc,
        "Governanca, Seguranca e Escalabilidade do INOVE_QUATAI",
        style="Subtitle",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=14,
    )
    add_para(
        doc,
        "Documento executivo-tecnico para endurecimento do sistema, preparacao multiempresa e reducao de risco operacional.",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=24,
        color=SLATE,
        size=11,
    )

    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(2.05)
    table.columns[1].width = Inches(3.75)
    set_table_borders(table, color="C9D6E3", size=8)

    data = [
        ("Projeto", "INOVE_QUATAI"),
        ("Supabase", "Projeto INOVEQUATAI - ref wboelthngddvkgrvwkbu"),
        ("GitHub", "Repositorio guuimaximo/INOVE_QUATAI"),
        ("Objetivo", "Elevar seguranca, governanca e prontidao para escalar o produto"),
    ]
    for idx, (k, v) in enumerate(data):
        c1, c2 = table.rows[idx].cells
        c1.text = k
        c2.text = v
        for cell in (c1, c2):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell, top=110, start=110, bottom=110, end=110)
        set_cell_shading(c1, "DCEAF7")
        set_cell_shading(c2, SOFT_GRAY)
        c1.paragraphs[0].runs[0].bold = True
        c1.paragraphs[0].runs[0].font.color.rgb = NAVY

    add_para(
        doc,
        "Data da revisao: 28/04/2026",
        style="Body Small",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        before=18,
        after=0,
        color=SLATE,
    )
    doc.add_section(WD_SECTION.NEW_PAGE)


def add_summary(doc):
    doc.add_heading("Resumo Executivo", level=1)
    add_para(
        doc,
        "O sistema ainda nao esta pronto para ser ofertado para outras empresas com seguranca maxima. "
        "Os dois bloqueios centrais sao o login legado no frontend usando usuarios_aprovadores.senha e a ausencia de um modelo multiempresa real com isolamento por empresa no banco e no codigo.",
        after=8,
    )
    add_para(
        doc,
        "Em paralelo, a schema publica do Supabase esta excessivamente exposta, o repositorio GitHub esta publico, "
        "o frontend preve uso de token do GitHub no navegador e varios buckets de storage estao publicos.",
        after=10,
    )

    box = doc.add_table(rows=1, cols=1)
    box.alignment = WD_TABLE_ALIGNMENT.CENTER
    box.autofit = False
    box.columns[0].width = Inches(6.0)
    set_table_borders(box, color="AFC4D9", size=8)
    cell = box.cell(0, 0)
    set_cell_shading(cell, SOFT_BLUE)
    set_cell_margins(cell, top=130, start=130, bottom=130, end=130)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(
        "Meta do plano: migrar o sistema de um produto single-tenant customizado para uma base segura, auditavel e pronta para escalar por empresa, com Quatai como empresa 046."
    )
    r.bold = True
    r.font.color.rgb = NAVY

    doc.add_paragraph()


def add_critical_findings(doc):
    doc.add_page_break()
    doc.add_heading("Prioridades Criticas", level=1)
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [0.72, 1.7, 2.15, 1.6]
    for idx, width in enumerate(widths):
        table.columns[idx].width = Inches(width)
    set_table_borders(table, color="D7E0EA", size=8)

    headers = ["Nivel", "Tema", "Situacao atual", "Acao imediata"]
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = text
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(cell, top=110, start=100, bottom=110, end=100)
        set_cell_shading(cell, "123B65")
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)

    rows = [
        ("P0", "Auth", "Login no frontend compara senha em usuarios_aprovadores", "Fechar login legado e migrar 100% para Supabase Auth"),
        ("P0", "Banco", "49 de 54 tabelas publicas sem RLS ativo", "Revisar grants e ativar RLS por dominio"),
        ("P0", "GitHub", "Frontend preve VITE_GITHUB_TOKEN no bundle", "Remover token do browser e mover automacoes para backend"),
        ("P1", "Storage", "Buckets sensiveis publicos com milhares de objetos", "Trocar para buckets privados e signed URLs"),
        ("P1", "Multiempresa", "Nao existe empresa_id ou tenant_id no schema publico", "Introduzir camada de tenancy e isolamento"),
        ("P2", "Codigo", "Arquivos duplicados, fluxos paralelos e rota backend no repo Vite", "Classificar e remover codigo obsoleto"),
    ]
    fills = {"P0": SOFT_RED, "P1": SOFT_AMBER, "P2": SOFT_GREEN}

    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            cells[idx].text = text
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cells[idx], top=100, start=100, bottom=100, end=100)
        set_cell_shading(cells[0], fills[row[0]])
        cells[0].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()


def add_scope_snapshot(doc):
    doc.add_heading("Escopo Revisado", level=1)
    add_bullets(
        doc,
        [
            "Supabase: tabelas publicas, colunas criticas, RLS, policies, views, funcoes SECURITY DEFINER e buckets de storage.",
            "GitHub: visibilidade do repositorio, PRs recentes, workflow de ETL e padrao de integracao no frontend.",
            "Frontend Vite: autenticacao, variaveis, storage publico, clientes Supabase, codigo duplicado e acoplamentos externos.",
            "Escalabilidade: prontidao para multiempresa, com foco no futuro uso da Quatai como empresa 046.",
        ],
    )


def add_phase_plan(doc):
    doc.add_heading("Plano de Acao por Fase", level=1)

    phases = [
        (
            "Fase 0 - Contencao imediata (0 a 7 dias)",
            [
                "Remover qualquer uso de VITE_GITHUB_TOKEN no frontend.",
                "Restringir ou remover auth_emails_view.",
                "Congelar novas funcionalidades dependentes de usuarios_aprovadores.senha.",
                "Mapear e reduzir grants de anon na schema public.",
                "Revisar publicacao atual e confirmar quais secrets estao expostos ao browser.",
            ],
        ),
        (
            "Fase 1 - Migracao de autenticacao (1 a 3 semanas)",
            [
                "Trocar login e reset para Supabase Auth.",
                "Usar profiles como fonte primaria de identidade e papeis.",
                "Encerrar autorizacao baseada em localStorage.",
                "Migrar usuarios_aprovadores para papel de tabela de transicao apenas.",
                "Planejar remocao da coluna senha ao final da migracao.",
            ],
        ),
        (
            "Fase 2 - Endurecimento do banco e storage (2 a 5 semanas)",
            [
                "Ativar RLS por dominio nas tabelas sensiveis.",
                "Reescrever policies permissivas e remover allow all.",
                "Revisar SECURITY DEFINER em views e funcoes.",
                "Migrar buckets sensiveis para privados.",
                "Limpar indices, constraints e objetos duplicados.",
            ],
        ),
        (
            "Fase 3 - Governanca de codigo e integracoes (3 a 6 semanas)",
            [
                "Classificar arquivos como ativo, legado, migrar ou remover.",
                "Remover codigo morto e fluxos paralelos de cadastro/login.",
                "Separar rotas backend do repositorio frontend.",
                "Corrigir o workflow ETL e documentar owner tecnico.",
                "Padronizar naming, ownership e criterios de deploy.",
            ],
        ),
        (
            "Fase 4 - Preparacao multiempresa (4 a 8 semanas)",
            [
                "Criar tabela empresas e chave empresa_id.",
                "Associar usuario autenticado a empresa e perfil.",
                "Adicionar empresa_id nas tabelas de negocio.",
                "Aplicar RLS com filtro por empresa.",
                "Fazer backfill da Quatai como empresa 046.",
            ],
        ),
    ]

    for title, bullets in phases:
        doc.add_heading(title, level=2)
        add_bullets(doc, bullets)


def add_action_matrix(doc):
    doc.add_page_break()
    doc.add_heading("Matriz de Execucao", level=1)
    table = doc.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [2.15, 1.0, 1.15, 1.0, 1.2]
    for idx, width in enumerate(widths):
        table.columns[idx].width = Inches(width)
    set_table_borders(table, color="D7E0EA", size=8)

    headers = ["Frente", "Prioridade", "Prazo", "Esforco", "Resultado esperado"]
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = text
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(cell, top=110, start=95, bottom=110, end=95)
        set_cell_shading(cell, "123B65")
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)

    rows = [
        ("Fechar login legado e reset inseguro", "P0", "Semana 1", "M", "Auth segura e pronta para endurecimento"),
        ("Revisar grants + ativar RLS nas tabelas criticas", "P0", "Semana 1-2", "G", "Reducao imediata da superficie de ataque"),
        ("Remover token GitHub do frontend", "P0", "Semana 1", "P", "Fim do risco de vazamento de token no browser"),
        ("Privatizar buckets sensiveis", "P1", "Semana 2-3", "M", "Anexos e relatorios protegidos"),
        ("Limpar codigo morto e fluxos paralelos", "P2", "Semana 3-4", "M", "Menos ambiguidade operacional"),
        ("Desenhar e aplicar empresa_id = 046", "P1", "Semana 4-6", "G", "Base pronta para multiempresa"),
    ]

    fills = {"P0": SOFT_RED, "P1": SOFT_AMBER, "P2": SOFT_GREEN}
    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            cells[idx].text = text
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cells[idx], top=95, start=95, bottom=95, end=95)
        set_cell_shading(cells[1], fills[row[1]])
        cells[1].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()


def add_multiempresa_section(doc):
    doc.add_heading("Diretriz de Multiempresa", level=1)
    add_para(
        doc,
        "Hoje o sistema esta mais proximo de um produto single-tenant customizado do que de uma plataforma multiempresa. "
        "Para escalar, a Quatai deve entrar formalmente no modelo como empresa 046, e nao apenas como convencao de negocio.",
        after=8,
    )

    add_bullets(
        doc,
        [
            "Criar tabela empresas com codigo_empresa, nome_fantasia, status e configuracoes de branding.",
            "Associar profiles a empresa_id e, se necessario, permitir multiplas empresas por usuario por tabela de vinculo.",
            "Adicionar empresa_id em tabelas de negocio e buckets/logica de storage.",
            "Aplicar RLS com filtro por empresa para leitura, escrita, update e delete.",
            "Retirar do browser qualquer integracao cross-base como supabaseBCNT.",
        ],
    )

    box = doc.add_table(rows=1, cols=1)
    box.alignment = WD_TABLE_ALIGNMENT.CENTER
    box.autofit = False
    box.columns[0].width = Inches(6.05)
    set_table_borders(box, color="E7B95A", size=8)
    cell = box.cell(0, 0)
    set_cell_shading(cell, SOFT_AMBER)
    set_cell_margins(cell, top=120, start=120, bottom=120, end=120)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(
        "Decisao recomendada: antes de onboardar qualquer nova empresa, concluir a modelagem de tenancy e fazer o backfill da Quatai como empresa 046."
    )
    run.bold = True
    run.font.color.rgb = AMBER
    doc.add_paragraph()


def add_obsolescence_section(doc):
    doc.add_heading("Codigo e Objetos Candidatos a Limpeza", level=1)
    add_para(
        doc,
        "Foram identificados arquivos duplicados e fluxos paralelos que aumentam ambiguidade e risco de manutencao. "
        "A limpeza deve ser feita de forma segura, marcando cada item como ativo, legado, migrar ou remover.",
        after=8,
    )
    items = [
        "src/AuthContext.jsx",
        "src/components/RequireAuth.jsx",
        "src/components/ProtectedRoute.jsx",
        "src/pages/auth/Cadastro.jsx",
        "src/pages/auth/Register.jsx",
        "src/pages/home/Landing.jsx",
        "src/pages/portal/PortalSistemas.jsx",
        "src/routes/premiacaoRoutes.js",
    ]
    add_bullets(doc, items)
    add_para(
        doc,
        "No banco, objetos sem referencia direta no frontend atual devem ser marcados como sem referencia no frontend, e nao removidos automaticamente. "
        "Eles podem estar em uso por ETL, dashboards externos, SQL manual ou processos internos.",
        style="Callout",
        before=6,
        after=0,
        color=SLATE,
    )


def add_next_deliverables(doc):
    doc.add_heading("Entregaveis Recomendados na Sequencia", level=1)
    add_numbered(
        doc,
        [
            "Matriz tabela -> owner -> sensibilidade -> RLS -> uso no codigo.",
            "Plano SQL de endurecimento do Supabase.",
            "Plano de migracao usuarios_aprovadores -> profiles/auth.",
            "Plano de tenancy com empresa_id = 046 para Quatai.",
            "Lista de remocao segura de codigo e objetos obsoletos.",
        ],
    )


def build_doc():
    doc = Document()
    style_document(doc)
    add_cover(doc)
    add_summary(doc)
    add_scope_snapshot(doc)
    add_critical_findings(doc)
    add_phase_plan(doc)
    add_action_matrix(doc)
    add_multiempresa_section(doc)
    doc.add_section(WD_SECTION.NEW_PAGE)
    add_obsolescence_section(doc)
    add_next_deliverables(doc)
    DOC_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(DOC_PATH)
    print(DOC_PATH)


if __name__ == "__main__":
    build_doc()
