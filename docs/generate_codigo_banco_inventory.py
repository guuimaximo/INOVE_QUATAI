from __future__ import annotations

import csv
import json
import re
from collections import defaultdict
from dataclasses import dataclass, field
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\Users\Guilh\AppData\Local\Temp\INOVE_QUATAI")
SRC = ROOT / "src"
DOCS = ROOT / "docs"

JSON_PATH = DOCS / "inventario_codigo_banco_inove.json"
CSV_PATH = DOCS / "inventario_codigo_banco_inove.csv"
MD_PATH = DOCS / "inventario_codigo_banco_inove.md"
DOCX_PATH = DOCS / "Inventario_Codigo_Banco_INOVE_QUATAI.docx"

NAVY = RGBColor(20, 62, 102)
SLATE = RGBColor(71, 85, 105)
TEAL = RGBColor(13, 148, 136)
RED = RGBColor(185, 28, 28)
AMBER = RGBColor(180, 83, 9)
SOFT_BLUE = "EAF2FB"
SOFT_GRAY = "F8FAFC"
SOFT_RED = "FDECEC"
SOFT_AMBER = "FEF3C7"
SOFT_GREEN = "EAF7EF"

SOURCE_EXTENSIONS = {".js", ".jsx", ".ts", ".tsx"}

FROM_RE = re.compile(r"(?P<client>\w+)\s*\.\s*from\(\s*['\"](?P<object>[^'\"]+)['\"]\s*\)")
RPC_RE = re.compile(r"(?P<client>\w+)\s*\.\s*rpc\(\s*['\"](?P<object>[^'\"]+)['\"]")
STORAGE_RE = re.compile(r"(?P<client>\w+)\s*\.\s*storage\s*\.\s*from\(\s*['\"](?P<object>[^'\"]+)['\"]\s*\)")
AUTH_RE = re.compile(r"(?P<client>\w+)\s*\.\s*auth\s*\.\s*(?P<method>\w+)\(")
CHANNEL_RE = re.compile(r"(?P<client>\w+)\s*\.\s*channel\(\s*['\"](?P<object>[^'\"]+)['\"]")
CREATE_CLIENT_RE = re.compile(r"createClient\(")
GITHUB_ENV_RE = re.compile(r"VITE_GITHUB_[A-Z_]+")
LOCALSTORAGE_RE = re.compile(r"localStorage\.(getItem|setItem|removeItem|clear)\(")
PUBLIC_STORAGE_URL_RE = re.compile(r"/storage/v1/object/public/(?P<bucket>[^/]+)/")
SUPABASE_URL_RE = re.compile(r"https://[a-z0-9-]+\.supabase\.co", re.IGNORECASE)

TABLE_OPERATIONS = ["select", "insert", "update", "upsert", "delete", "eq", "in", "single", "maybeSingle"]
STORAGE_OPERATIONS = ["upload", "download", "remove", "list", "createSignedUrl", "createSignedUrls", "getPublicUrl"]
RISK_PATTERNS = [
    ("github_token_frontend", re.compile(r"VITE_GITHUB_TOKEN")),
    ("service_role_in_frontend", re.compile(r"SUPABASE_SERVICE_ROLE_KEY")),
    ("legacy_password_query", re.compile(r"\.eq\(\s*['\"]senha['\"]")),
    ("cross_base_client", re.compile(r"supabaseBCNT")),
    ("hardcoded_supabase_url", SUPABASE_URL_RE),
]


@dataclass
class Usage:
    object_name: str
    category: str
    client: str
    line: int
    operations: list[str] = field(default_factory=list)


@dataclass
class FileInventory:
    path: str
    module: str
    clients: set[str] = field(default_factory=set)
    tables_views: list[Usage] = field(default_factory=list)
    buckets: list[Usage] = field(default_factory=list)
    rpcs: list[Usage] = field(default_factory=list)
    auth_methods: list[Usage] = field(default_factory=list)
    channels: list[Usage] = field(default_factory=list)
    flags: set[str] = field(default_factory=set)
    notes: list[str] = field(default_factory=list)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=90, start=90, bottom=90, end=90) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="D7E0EA", size=6) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(size))
        element.set(qn("w:color"), color)


def classify_module(path: Path) -> str:
    rel = path.relative_to(ROOT).as_posix()
    if rel.startswith("src/pages/"):
        return "page"
    if rel.startswith("src/components/"):
        return "component"
    if rel.startswith("src/context/") or rel in {"src/AuthContext.jsx"}:
        return "auth_context"
    if rel.startswith("src/routes/"):
        return "route"
    if rel.startswith("src/utils/"):
        return "util"
    if path.name in {"supabase.js", "supabaseBCNT.js"}:
        return "supabase_client"
    return "app_core"


def dedupe_usages(usages: Iterable[Usage]) -> list[Usage]:
    seen = set()
    deduped: list[Usage] = []
    for usage in usages:
        key = (usage.object_name, usage.category, usage.client, usage.line, tuple(usage.operations))
        if key not in seen:
            seen.add(key)
            deduped.append(usage)
    return deduped


def scan_operations(text: str, start: int, candidates: list[str]) -> list[str]:
    window = text[start : start + 400]
    semicolon = window.find(";")
    if semicolon != -1:
        window = window[:semicolon]
    ops = [op for op in candidates if f".{op}(" in window]
    return ops or ["unknown"]


def line_number_from_offset(text: str, offset: int) -> int:
    return text.count("\n", 0, offset) + 1


def scan_file(path: Path) -> FileInventory:
    text = path.read_text(encoding="utf-8", errors="ignore")
    rel = path.relative_to(ROOT).as_posix()
    inventory = FileInventory(path=rel, module=classify_module(path))

    for pattern_name, pattern in RISK_PATTERNS:
        if pattern.search(text):
            inventory.flags.add(pattern_name)

    if LOCALSTORAGE_RE.search(text):
        inventory.flags.add("localstorage_session")

    if GITHUB_ENV_RE.search(text):
        inventory.flags.add("github_env_usage")

    if CREATE_CLIENT_RE.search(text):
        inventory.flags.add("create_client_local")
        inventory.notes.append("Arquivo instancia cliente Supabase localmente.")

    for match in FROM_RE.finditer(text):
        inventory.clients.add(match.group("client"))
        inventory.tables_views.append(
            Usage(
                object_name=match.group("object"),
                category="table_or_view",
                client=match.group("client"),
                line=line_number_from_offset(text, match.start()),
                operations=scan_operations(text, match.start(), TABLE_OPERATIONS),
            )
        )

    for match in RPC_RE.finditer(text):
        inventory.clients.add(match.group("client"))
        inventory.rpcs.append(
            Usage(
                object_name=match.group("object"),
                category="rpc",
                client=match.group("client"),
                line=line_number_from_offset(text, match.start()),
                operations=scan_operations(text, match.start(), TABLE_OPERATIONS),
            )
        )

    for match in STORAGE_RE.finditer(text):
        inventory.clients.add(match.group("client"))
        inventory.buckets.append(
            Usage(
                object_name=match.group("object"),
                category="bucket",
                client=match.group("client"),
                line=line_number_from_offset(text, match.start()),
                operations=scan_operations(text, match.start(), STORAGE_OPERATIONS),
            )
        )

    for match in AUTH_RE.finditer(text):
        inventory.clients.add(match.group("client"))
        inventory.auth_methods.append(
            Usage(
                object_name=match.group("method"),
                category="auth_method",
                client=match.group("client"),
                line=line_number_from_offset(text, match.start()),
                operations=[match.group("method")],
            )
        )

    for match in CHANNEL_RE.finditer(text):
        inventory.clients.add(match.group("client"))
        inventory.channels.append(
            Usage(
                object_name=match.group("object"),
                category="channel",
                client=match.group("client"),
                line=line_number_from_offset(text, match.start()),
                operations=["channel"],
            )
        )

    for match in PUBLIC_STORAGE_URL_RE.finditer(text):
        inventory.flags.add("hardcoded_public_storage_url")
        inventory.buckets.append(
            Usage(
                object_name=match.group("bucket"),
                category="bucket_public_url",
                client="url_literal",
                line=line_number_from_offset(text, match.start()),
                operations=["public_url"],
            )
        )

    inventory.tables_views = dedupe_usages(inventory.tables_views)
    inventory.buckets = dedupe_usages(inventory.buckets)
    inventory.rpcs = dedupe_usages(inventory.rpcs)
    inventory.auth_methods = dedupe_usages(inventory.auth_methods)
    inventory.channels = dedupe_usages(inventory.channels)
    inventory.notes = sorted(set(inventory.notes))
    return inventory


def scan_repo() -> list[FileInventory]:
    files = sorted(
        path
        for path in SRC.rglob("*")
        if path.is_file() and path.suffix in SOURCE_EXTENSIONS
    )
    inventories = [scan_file(path) for path in files]
    return [
        inv
        for inv in inventories
        if inv.tables_views or inv.buckets or inv.rpcs or inv.auth_methods or inv.channels or inv.flags
    ]


def build_object_index(inventories: list[FileInventory]) -> dict[str, dict]:
    index: dict[str, dict] = {}

    def ensure(name: str, category: str) -> dict:
        key = f"{category}::{name}"
        return index.setdefault(key, {"name": name, "category": category, "references": []})

    for inv in inventories:
        for collection in (inv.tables_views, inv.buckets, inv.rpcs, inv.auth_methods, inv.channels):
            for usage in collection:
                bucket = ensure(usage.object_name, usage.category)
                bucket["references"].append(
                    {
                        "file": inv.path,
                        "line": usage.line,
                        "client": usage.client,
                        "operations": usage.operations,
                    }
                )
    return dict(sorted(index.items(), key=lambda item: (item[1]["category"], item[1]["name"])))


def build_summary(inventories: list[FileInventory], object_index: dict[str, dict]) -> dict:
    files_by_module = defaultdict(int)
    flags = defaultdict(list)
    clients = defaultdict(int)
    for inv in inventories:
        files_by_module[inv.module] += 1
        for flag in sorted(inv.flags):
            flags[flag].append(inv.path)
        for client in inv.clients:
            clients[client] += 1

    distinct_categories = {
        data["category"]
        for data in object_index.values()
    }

    return {
        "files_with_db_touchpoints": len(inventories),
        "distinct_objects": len(object_index),
        "distinct_tables_views": sum(1 for item in object_index.values() if item["category"] == "table_or_view"),
        "distinct_buckets": sum(1 for item in object_index.values() if item["category"].startswith("bucket")),
        "distinct_rpcs": sum(1 for item in object_index.values() if item["category"] == "rpc"),
        "distinct_auth_methods": sum(1 for item in object_index.values() if item["category"] == "auth_method"),
        "distinct_object_categories": sorted(distinct_categories),
        "files_by_module": dict(sorted(files_by_module.items())),
        "flags": dict(sorted(flags.items())),
        "clients": dict(sorted(clients.items())),
    }


def write_json(inventories: list[FileInventory], object_index: dict[str, dict], summary: dict) -> None:
    payload = {
        "summary": summary,
        "files": [
            {
                "path": inv.path,
                "module": inv.module,
                "clients": sorted(inv.clients),
                "tables_views": [usage.__dict__ for usage in inv.tables_views],
                "buckets": [usage.__dict__ for usage in inv.buckets],
                "rpcs": [usage.__dict__ for usage in inv.rpcs],
                "auth_methods": [usage.__dict__ for usage in inv.auth_methods],
                "channels": [usage.__dict__ for usage in inv.channels],
                "flags": sorted(inv.flags),
                "notes": inv.notes,
            }
            for inv in inventories
        ],
        "objects": object_index,
    }
    JSON_PATH.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")


def write_csv(inventories: list[FileInventory]) -> None:
    rows = []
    for inv in inventories:
        for label, collection in (
            ("table_or_view", inv.tables_views),
            ("bucket", inv.buckets),
            ("rpc", inv.rpcs),
            ("auth_method", inv.auth_methods),
            ("channel", inv.channels),
        ):
            for usage in collection:
                rows.append(
                    {
                        "file": inv.path,
                        "module": inv.module,
                        "client": usage.client,
                        "category": label,
                        "object": usage.object_name,
                        "line": usage.line,
                        "operations": ", ".join(usage.operations),
                        "flags": ", ".join(sorted(inv.flags)),
                    }
                )
        if inv.flags and not any((inv.tables_views, inv.buckets, inv.rpcs, inv.auth_methods, inv.channels)):
            rows.append(
                {
                    "file": inv.path,
                    "module": inv.module,
                    "client": "",
                    "category": "flag_only",
                    "object": "",
                    "line": "",
                    "operations": "",
                    "flags": ", ".join(sorted(inv.flags)),
                }
            )

    with CSV_PATH.open("w", newline="", encoding="utf-8") as fh:
        writer = csv.DictWriter(
            fh,
            fieldnames=["file", "module", "client", "category", "object", "line", "operations", "flags"],
        )
        writer.writeheader()
        writer.writerows(rows)


def join_usage(usages: list[Usage]) -> str:
    if not usages:
        return "-"
    return ", ".join(f"{item.object_name} (l.{item.line})" for item in usages)


def write_markdown(inventories: list[FileInventory], object_index: dict[str, dict], summary: dict) -> None:
    lines: list[str] = []
    lines.append("# Inventario Codigo x Banco do INOVE_QUATAI")
    lines.append("")
    lines.append("Documento gerado automaticamente a partir do codigo em `src/` para apoiar handoff tecnico, governanca e saneamento.")
    lines.append("")
    lines.append("## Resumo")
    lines.append("")
    lines.append(f"- Arquivos com touchpoints de banco/storage/auth: {summary['files_with_db_touchpoints']}")
    lines.append(f"- Objetos distintos mapeados: {summary['distinct_objects']}")
    lines.append(f"- Tabelas/views distintas: {summary['distinct_tables_views']}")
    lines.append(f"- Buckets distintos: {summary['distinct_buckets']}")
    lines.append(f"- RPCs distintas: {summary['distinct_rpcs']}")
    lines.append(f"- Metodos de auth distintos: {summary['distinct_auth_methods']}")
    lines.append("")
    lines.append("## Flags criticas")
    lines.append("")
    for flag, files in summary["flags"].items():
        lines.append(f"- `{flag}`: {len(files)} arquivo(s)")
        for file in files[:12]:
            lines.append(f"  - `{file}`")
    lines.append("")
    lines.append("## Matriz por arquivo")
    lines.append("")
    lines.append("| Arquivo | Modulo | Clientes | Tabelas/Views | Buckets | RPC/Auth | Flags |")
    lines.append("| --- | --- | --- | --- | --- | --- | --- |")
    for inv in inventories:
        rpc_auth = join_usage(inv.rpcs + inv.auth_methods)
        lines.append(
            f"| `{inv.path}` | `{inv.module}` | `{', '.join(sorted(inv.clients)) or '-'}` | "
            f"{join_usage(inv.tables_views)} | {join_usage(inv.buckets)} | {rpc_auth} | "
            f"`{', '.join(sorted(inv.flags)) or '-'}` |"
        )
    lines.append("")
    lines.append("## Matriz por objeto")
    lines.append("")
    grouped = defaultdict(list)
    for data in object_index.values():
        grouped[data["category"]].append(data)

    category_titles = {
        "table_or_view": "Tabelas e Views",
        "bucket": "Buckets",
        "bucket_public_url": "Buckets via URL publica",
        "rpc": "RPCs",
        "auth_method": "Auth",
        "channel": "Realtime/Channels",
    }

    for category in ("table_or_view", "bucket", "bucket_public_url", "rpc", "auth_method", "channel"):
        objects = grouped.get(category)
        if not objects:
            continue
        lines.append(f"### {category_titles.get(category, category)}")
        lines.append("")
        for data in objects:
            lines.append(f"- `{data['name']}`")
            for ref in sorted(data["references"], key=lambda item: (item["file"], item["line"])):
                ops = ", ".join(ref["operations"])
                lines.append(
                    f"  - `{ref['file']}:{ref['line']}` via `{ref['client']}` com operacoes `{ops}`"
                )
            lines.append("")
    MD_PATH.write_text("\n".join(lines), encoding="utf-8")


def style_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.82)
    section.right_margin = Inches(0.82)

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10)
    normal.font.color.rgb = RGBColor(31, 41, 55)

    title = doc.styles["Title"]
    title.font.name = "Aptos Display"
    title.font.size = Pt(22)
    title.font.bold = True
    title.font.color.rgb = NAVY

    subtitle = doc.styles["Subtitle"]
    subtitle.font.name = "Aptos"
    subtitle.font.size = Pt(11)
    subtitle.font.color.rgb = SLATE

    for name, size, color in (
        ("Heading 1", 15, NAVY),
        ("Heading 2", 11.5, TEAL),
        ("Heading 3", 10.5, SLATE),
    ):
        style = doc.styles[name]
        style.font.name = "Aptos"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color

    if "Body Small" not in doc.styles:
        style = doc.styles.add_style("Body Small", WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = normal
        style.font.name = "Aptos"
        style.font.size = Pt(9)


def add_paragraph(doc: Document, text: str, style=None, before=0, after=6, bold=False, color=None, align=None, size=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    if size is not None:
        run.font.size = Pt(size)
    return p


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        p.add_run(item)


def add_cover(doc: Document, summary: dict) -> None:
    add_paragraph(doc, "Inventario Codigo x Banco", style="Title", align=WD_ALIGN_PARAGRAPH.CENTER, before=64, after=2)
    add_paragraph(
        doc,
        "INOVE_QUATAI - Mapeamento de tabelas, views, buckets e fluxos Supabase no frontend",
        style="Subtitle",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=14,
    )
    add_paragraph(
        doc,
        "Documento de handoff tecnico para desenvolvimento, saneamento e governanca da ferramenta.",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        color=SLATE,
        size=11,
        after=18,
    )

    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(2.1)
    table.columns[1].width = Inches(3.9)
    set_table_borders(table, color="C9D6E3", size=8)
    pairs = [
        ("Projeto", "INOVE_QUATAI"),
        ("Escopo", "Frontend Vite em src/ com clientes supabase e supabaseBCNT"),
        ("Objetos mapeados", str(summary["distinct_objects"])),
        ("Arquivos com touchpoints", str(summary["files_with_db_touchpoints"])),
    ]
    for idx, (left, right) in enumerate(pairs):
        c1, c2 = table.rows[idx].cells
        c1.text = left
        c2.text = right
        for cell in (c1, c2):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell, 100, 100, 100, 100)
        set_cell_shading(c1, "DCEAF7")
        set_cell_shading(c2, SOFT_GRAY)
        c1.paragraphs[0].runs[0].bold = True
        c1.paragraphs[0].runs[0].font.color.rgb = NAVY

    add_paragraph(doc, "Data da geracao: 28/04/2026", style="Body Small", align=WD_ALIGN_PARAGRAPH.CENTER, before=16, after=0, color=SLATE)
    doc.add_page_break()


def add_summary_section(doc: Document, summary: dict) -> None:
    doc.add_heading("Resumo Tecnico", level=1)
    add_paragraph(
        doc,
        "Este inventario cruza o codigo do frontend com os objetos de banco e storage para acelerar onboarding de dev, revisao de seguranca e saneamento de acoplamentos.",
        after=8,
    )
    add_bullets(
        doc,
        [
            f"{summary['files_with_db_touchpoints']} arquivos fazem chamadas diretas a banco, storage ou auth.",
            f"{summary['distinct_tables_views']} tabelas/views distintas aparecem nas consultas do frontend.",
            f"{summary['distinct_buckets']} buckets distintos aparecem por SDK ou URL publica.",
            f"{summary['distinct_auth_methods']} metodos de auth distintos aparecem no codigo.",
        ],
    )

    box = doc.add_table(rows=1, cols=1)
    box.alignment = WD_TABLE_ALIGNMENT.CENTER
    box.autofit = False
    box.columns[0].width = Inches(6.1)
    set_table_borders(box, color="AFC4D9", size=8)
    cell = box.cell(0, 0)
    set_cell_shading(cell, SOFT_BLUE)
    set_cell_margins(cell, 120, 120, 120, 120)
    run = cell.paragraphs[0].add_run(
        "Leitura recomendada para o dev: 1) flags criticas, 2) arquitetura de clientes, 3) matriz por arquivo, 4) matriz por objeto."
    )
    run.bold = True
    run.font.color.rgb = NAVY


def add_flags_section(doc: Document, summary: dict) -> None:
    doc.add_heading("Flags Criticas", level=1)
    descriptions = {
        "github_env_usage": "Arquivo consome variaveis VITE_GITHUB_* no frontend.",
        "github_token_frontend": "Arquivo referencia VITE_GITHUB_TOKEN diretamente.",
        "service_role_in_frontend": "Arquivo referencia service role em repositorio frontend.",
        "legacy_password_query": "Arquivo toca no fluxo legado que consulta usuarios_aprovadores.senha.",
        "cross_base_client": "Arquivo usa cliente supabaseBCNT e expande a superficie multi-base.",
        "hardcoded_public_storage_url": "Arquivo monta URL publica de storage manualmente.",
        "hardcoded_supabase_url": "Arquivo contem URL Supabase hardcoded.",
        "localstorage_session": "Arquivo depende de sessao/estado via localStorage.",
        "create_client_local": "Arquivo instancia cliente Supabase localmente.",
    }

    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(1.85)
    table.columns[1].width = Inches(2.45)
    table.columns[2].width = Inches(2.0)
    set_table_borders(table, color="D7E0EA", size=8)
    for cell, text in zip(table.rows[0].cells, ["Flag", "Significado", "Arquivos"]):
        cell.text = text
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(cell, 100, 100, 100, 100)
        set_cell_shading(cell, "123B65")
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    fills = {"github_token_frontend": SOFT_RED, "service_role_in_frontend": SOFT_RED, "legacy_password_query": SOFT_RED}
    for flag, files in summary["flags"].items():
        row = table.add_row().cells
        row[0].text = flag
        row[1].text = descriptions.get(flag, "-")
        row[2].text = str(len(files))
        for cell in row:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell, 90, 90, 90, 90)
        set_cell_shading(row[0], fills.get(flag, SOFT_AMBER if len(files) > 1 else SOFT_GREEN))


def compact_list(values: list[str], max_items: int = 5) -> str:
    if not values:
        return "-"
    display = values[:max_items]
    suffix = f" +{len(values) - max_items}" if len(values) > max_items else ""
    return ", ".join(display) + suffix


def add_file_matrix(doc: Document, inventories: list[FileInventory]) -> None:
    doc.add_page_break()
    doc.add_heading("Matriz por Arquivo", level=1)
    add_paragraph(
        doc,
        "A tabela abaixo mostra, para cada arquivo com touchpoint de banco, quais clientes, objetos e sinais de risco aparecem no codigo.",
        after=8,
    )

    table = doc.add_table(rows=1, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [2.25, 0.95, 0.9, 1.6, 1.35, 1.05]
    for idx, width in enumerate(widths):
        table.columns[idx].width = Inches(width)
    set_table_borders(table, color="D7E0EA", size=6)

    headers = ["Arquivo", "Modulo", "Cliente", "Tabelas/Views", "Buckets/RPC", "Flags"]
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = text
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(cell, 90, 90, 90, 90)
        set_cell_shading(cell, "123B65")
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    for inv in inventories:
        row = table.add_row().cells
        row[0].text = inv.path.replace("src/", "", 1)
        row[1].text = inv.module
        row[2].text = compact_list(sorted(inv.clients), 3)
        row[3].text = compact_list([f"{u.object_name} (l.{u.line})" for u in inv.tables_views], 4)
        row[4].text = compact_list(
            [f"{u.object_name} (l.{u.line})" for u in (inv.buckets + inv.rpcs + inv.auth_methods)],
            4,
        )
        row[5].text = compact_list(sorted(inv.flags), 3)
        for cell in row:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell, 78, 78, 78, 78)
        if any(flag in inv.flags for flag in {"github_token_frontend", "service_role_in_frontend", "legacy_password_query"}):
            set_cell_shading(row[5], SOFT_RED)
        elif inv.flags:
            set_cell_shading(row[5], SOFT_AMBER)


def add_object_sections(doc: Document, object_index: dict[str, dict]) -> None:
    doc.add_page_break()
    doc.add_heading("Matriz por Objeto", level=1)
    category_titles = {
        "table_or_view": "Tabelas e Views",
        "bucket": "Buckets",
        "bucket_public_url": "Buckets via URL publica",
        "rpc": "RPCs",
        "auth_method": "Auth",
        "channel": "Realtime/Channels",
    }

    grouped = defaultdict(list)
    for data in object_index.values():
        grouped[data["category"]].append(data)

    for category in ("table_or_view", "bucket", "bucket_public_url", "rpc", "auth_method", "channel"):
        objects = grouped.get(category)
        if not objects:
            continue
        doc.add_heading(category_titles.get(category, category), level=2)
        for data in objects:
            doc.add_heading(data["name"], level=3)
            for ref in sorted(data["references"], key=lambda item: (item["file"], item["line"])):
                ops = ", ".join(ref["operations"])
                add_paragraph(
                    doc,
                    f"{ref['file']}:{ref['line']} via {ref['client']} | operacoes: {ops}",
                    style="Body Small",
                    after=1,
                )


def add_handoff_section(doc: Document, summary: dict) -> None:
    doc.add_page_break()
    doc.add_heading("Diretrizes de Handoff para DEV", level=1)
    add_bullets(
        doc,
        [
            "Usar este documento junto do plano de acao e do inventario CSV para priorizar limpeza por modulo.",
            "Comecar pelos arquivos com flags vermelhas: github token, service role, senha legada e cross-base.",
            "Separar logo no inicio os clientes supabase e supabaseBCNT, definindo o que continua e o que sai do browser.",
            "Antes de remover qualquer tabela sem referencia no frontend, validar ETL, dashboards e automacoes externas.",
            "Depois de cada refactor, regenerar este inventario para confirmar reducao real de acoplamentos.",
        ],
    )
    add_paragraph(
        doc,
        f"Clientes detectados no codigo: {', '.join(summary['clients'].keys()) or '-'}",
        before=8,
        after=0,
        bold=True,
        color=SLATE,
    )


def write_docx(inventories: list[FileInventory], object_index: dict[str, dict], summary: dict) -> None:
    doc = Document()
    style_doc(doc)
    add_cover(doc, summary)
    add_summary_section(doc, summary)
    add_flags_section(doc, summary)
    add_file_matrix(doc, inventories)
    add_object_sections(doc, object_index)
    add_handoff_section(doc, summary)
    DOCX_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(DOCX_PATH)


def main() -> None:
    DOCS.mkdir(parents=True, exist_ok=True)
    inventories = scan_repo()
    object_index = build_object_index(inventories)
    summary = build_summary(inventories, object_index)
    write_json(inventories, object_index, summary)
    write_csv(inventories)
    write_markdown(inventories, object_index, summary)
    write_docx(inventories, object_index, summary)
    print(JSON_PATH)
    print(CSV_PATH)
    print(MD_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    main()
